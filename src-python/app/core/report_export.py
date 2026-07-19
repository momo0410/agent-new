"""Multi-format report export from one normalized, integrity-bound document."""
from __future__ import annotations

import hashlib
import html
import json
import re
import textwrap
from collections.abc import Iterable
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Literal

ReportFormat = Literal["md", "html", "json", "pdf", "docx"]

_SECRET_KEYS = re.compile(r"(?:password|secret|token|api[_-]?key|private[_-]?key|cookie)", re.I)
_MISSING = "[MISSING]"
_HIGH_RISK_SEVERITIES = {"high", "critical", "严重", "高危"}
_HIGH_RISK_STATUSES = {
    "confirmed",
    "vulnerable",
    "exploited",
    "vulnerability_confirmed",
    "verified",
    "identity_confirmed",
}
_HIGH_RISK_FIELDS = (
    "why_suspected",
    "verification_method",
    "proof_evidence",
    "impact",
    "reproduction",
    "remediation",
)


def _scrub(value: Any) -> Any:
    if isinstance(value, dict):
        return {
            str(key): "[REDACTED]" if _SECRET_KEYS.search(str(key)) else _scrub(item)
            for key, item in value.items()
        }
    if isinstance(value, list):
        return [_scrub(item) for item in value]
    if isinstance(value, tuple):
        return [_scrub(item) for item in value]
    return value


def _first_value(item: dict[str, Any], *keys: str) -> Any:
    for key in keys:
        value = item.get(key)
        if value not in (None, "", [], {}):
            return value
    return None


def _text_or_missing(value: Any) -> Any:
    if value in (None, "", [], {}):
        return _MISSING
    if isinstance(value, str):
        return value.strip() or _MISSING
    return value


def _refs(item: dict[str, Any]) -> list[str]:
    values: list[str] = []
    for key in ("evidence_refs", "evidence_ids", "source_refs", "action_refs"):
        raw = item.get(key, [])
        if isinstance(raw, (list, tuple, set)):
            values.extend(str(value) for value in raw if str(value).strip())
        elif raw not in (None, ""):
            values.append(str(raw))
    for key in ("evidence_id", "action_id", "finding_id", "endpoint_id", "observation_id"):
        value = item.get(key)
        if value not in (None, ""):
            values.append(str(value))
    return list(dict.fromkeys(values))


def _source_type(item: dict[str, Any], default: str) -> str:
    provenance = item.get("provenance")
    explicit = item.get("source_type")
    if not explicit and isinstance(provenance, dict):
        explicit = provenance.get("source_type") or provenance.get("origin")
    return str(explicit or default).strip()[:80]


def _is_high_risk(item: dict[str, Any]) -> bool:
    severity = str(item.get("severity", "")).strip().lower()
    status = str(item.get("status", "")).strip().lower().replace(" ", "_")
    try:
        score = float(item.get("score", item.get("cvss_score", 0)) or 0)
    except (TypeError, ValueError):
        score = 0.0
    return severity in _HIGH_RISK_SEVERITIES or status in _HIGH_RISK_STATUSES or score >= 70


def _normalize_finding(item: dict[str, Any], *, source_group: str, index: int) -> dict[str, Any]:
    """Attach an explicit evidence-completeness envelope without inventing facts."""
    if source_group == "web_findings":
        default_source = "deterministic_web_rule"
    elif source_group == "vulnerabilities":
        default_source = "state_vulnerability"
    else:
        default_source = "tool_observation"
    refs = _refs(item)
    proof = _first_value(item, "proof_evidence", "proof", "evidence")
    if proof in (None, "", [], {}) and refs:
        proof = refs
    normalized = {
        "finding_id": str(
            _first_value(item, "finding_id", "id", "vulnerability_id")
            or f"{source_group}:{index + 1}"
        ),
        "source_group": source_group,
        "source_type": _source_type(item, default_source),
        "claim_type": str(item.get("claim_type") or item.get("classification") or "observed_record")[:60],
        "target": _text_or_missing(_first_value(item, "target", "ip", "host", "url")),
        "title": _text_or_missing(_first_value(item, "title", "name", "rule_id", "category")),
        "severity": _text_or_missing(item.get("severity")),
        "status": _text_or_missing(item.get("status")),
        "evidence_refs": refs,
        "why_suspected": _text_or_missing(_first_value(item, "why_suspected", "suspicion_reason", "reason")),
        "verification_method": _text_or_missing(
            _first_value(item, "verification_method", "validation_method", "verification", "judge_rule")
        ),
        "proof_evidence": _text_or_missing(proof),
        "impact": _text_or_missing(_first_value(item, "impact", "business_impact", "impact_analysis")),
        "reproduction": _text_or_missing(
            _first_value(item, "reproduction", "reproduction_steps", "reproduce", "steps")
        ),
        "remediation": _text_or_missing(_first_value(item, "remediation", "recommendation", "fix")),
        "detection_points": _text_or_missing(_first_value(item, "detection_points", "detection", "logs")),
        "log_sources": _text_or_missing(_first_value(item, "log_sources", "log_source")),
        "remediation_priority": _text_or_missing(_first_value(item, "remediation_priority", "priority")),
        "source_refs": refs,
        "source_record": _scrub(dict(item)),
    }
    if source_group == "web_findings" and normalized["verification_method"] == _MISSING:
        rule_id = item.get("rule_id")
        if rule_id:
            normalized["verification_method"] = f"WebRuleEngine:{rule_id}"
    missing = [field for field in _HIGH_RISK_FIELDS if normalized[field] == _MISSING]
    normalized["missing_fields"] = missing
    normalized["evidence_complete"] = not missing if _is_high_risk(item) else bool(refs)
    normalized["risk_class"] = "high" if _is_high_risk(item) else "standard"
    return normalized


def _canonical_hash(value: Any) -> str:
    encoded = json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"), default=str)
    return hashlib.sha256(encoded.encode("utf-8")).hexdigest()


def _build_provenance(data: dict[str, Any]) -> dict[str, Any]:
    scope = data.get("scope_contract") or data.get("scope") or {}
    if not isinstance(scope, dict):
        scope = {}
    scope_hash = str(
        data.get("scope_contract_hash")
        or data.get("scope_hash")
        or scope.get("canonical_hash")
        or (_canonical_hash(scope) if scope else "")
    )
    event_integrity = (
        data.get("event_integrity")
        or data.get("event_manifest")
        or data.get("event_store_manifest")
        or {}
    )
    if not isinstance(event_integrity, dict):
        event_integrity = {"last_hash": str(event_integrity)}
    browser_traces = data.get("browser_traces") or {}
    if isinstance(browser_traces, dict):
        browser_trace_ids = sorted(str(key) for key in browser_traces)
    elif isinstance(browser_traces, list):
        browser_trace_ids = sorted(
            str(item.get("trace_id")) for item in browser_traces if isinstance(item, dict) and item.get("trace_id")
        )
    else:
        browser_trace_ids = []
    model_calls = data.get("model_calls") or data.get("model_gateway_calls") or data.get("llm_calls") or []
    if not isinstance(model_calls, list):
        model_calls = [model_calls]
    return {
        "event_count": data.get("event_count", 0),
        "event_schema_version": data.get("event_schema_version", "event.v1"),
        "event_integrity": _scrub(event_integrity),
        "event_chain_hash": str(event_integrity.get("last_hash") or data.get("event_chain_hash", "")),
        "evidence_rule_version": str(data.get("evidence_rule_version", "evidence.v1")),
        "source": "state-and-event-store",
        "environment_fingerprint": _scrub(data.get("environment_fingerprint", {})),
        "runtime_snapshot": _scrub(data.get("runtime_snapshot", {})),
        "tool_versions": _scrub(
            (data.get("runtime_snapshot") or {}).get("tools", {})
            if isinstance(data.get("runtime_snapshot"), dict)
            else (data.get("versions") or {}).get("tools", {})
        ),
        "scope_contract_id": str(scope.get("scope_id") or data.get("scope_id", "")),
        "scope_contract_hash": scope_hash,
        "scope_revision": scope.get("revision", ""),
        "policy_version": str(scope.get("policy_version") or data.get("policy_version", "")),
        "policy_template": {
            "id": str(scope.get("policy_template_id", "")),
            "version": str(scope.get("policy_template_version", "")),
            "hash": str(scope.get("policy_template_hash", "")),
        },
        "scope_token_id": str(data.get("scope_token_id", "")),
        "autonomy_mode": str(data.get("autonomy_mode") or scope.get("autonomy_mode", "")),
        "autonomy_history": _scrub(list(data.get("autonomy_history", []) or [])[-100:]),
        "action_limit": str(data.get("action_limit", "")),
        "action_limit_history": _scrub(list(data.get("action_limit_history", []) or [])[-100:]),
        "tool_fallbacks": _scrub(list(data.get("tool_fallbacks", []) or [])[-200:]),
        "model_calls": _scrub(model_calls[-200:]),
        "token_usage": _scrub(data.get("token_usage", {})),
        "browser_trace_ids": browser_trace_ids,
        "observation_ids": sorted(
            str(item.get("observation_id"))
            for item in (data.get("observations") or data.get("web_observations") or [])
            if isinstance(item, dict) and item.get("observation_id")
        ),
        "versions": _scrub(data.get("versions", {})),
    }


def build_report_document(data: dict[str, Any], *, task_id: str) -> dict[str, Any]:
    """Build the canonical report payload consumed by every renderer."""
    findings = [item for item in data.get("findings", []) if isinstance(item, dict)]
    vulnerabilities = [item for item in data.get("vulnerabilities", []) if isinstance(item, dict)]
    web_findings = [item for item in data.get("web_findings", []) if isinstance(item, dict)]
    evidence = [item for item in data.get("canonical_evidence", []) if isinstance(item, dict)]
    actions = [item for item in data.get("actions_taken", []) if isinstance(item, dict)]
    technical_findings = [
        _normalize_finding(item, source_group=group, index=index)
        for group, records in (
            ("findings", findings),
            ("vulnerabilities", vulnerabilities),
            ("web_findings", web_findings),
        )
        for index, item in enumerate(records)
    ]
    failures = [
        {
            "action_id": item.get("id", ""),
            "tool": item.get("tool", ""),
            "failure_type": item.get("failure_type", "unknown"),
            "reason": item.get("error") or item.get("result") or "",
            "retry_condition": item.get("next_step", "改变前置条件、工具或路径后再评估"),
        }
        for item in actions
        if item.get("error") or str(item.get("status", "")).lower() in {"failed", "timeout", "error"}
    ]
    attack_path = [
        {
            "step": index,
            "action_id": item.get("id", ""),
            "tool": item.get("tool", ""),
            "purpose": item.get("purpose") or item.get("choice_reason", ""),
            "status": item.get("status", ""),
            "evidence_refs": [item.get("evidence_id")] if item.get("evidence_id") else [],
        }
        for index, item in enumerate(actions, 1)
    ]
    document = {
        "schema_version": "report-document.v1",
        "task_id": str(task_id),
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "management_summary": {
            "target_count": len(data.get("targets", [])),
            "finding_count": len(findings),
            "vulnerability_count": len(vulnerabilities),
            "web_finding_count": len(web_findings),
            "technical_finding_count": len(technical_findings),
            "high_risk_count": sum(item["risk_class"] == "high" for item in technical_findings),
            "evidence_count": len(evidence),
            "mission_status": data.get("mission_status", ""),
        },
        "technical_findings": _scrub(technical_findings),
        "attack_path": _scrub(attack_path),
        "evidence": _scrub(evidence),
        "failure_path": _scrub(failures),
        "coverage": _scrub({
            "targets": data.get("targets", []),
            "attack_surfaces": data.get("attack_surfaces", []),
            "asset_graph": data.get("asset_graph", {}),
            "web_findings": web_findings,
            "web_observations": data.get("web_observations", []),
            "web_sites": data.get("web_sites", {}),
            "browser_trace_ids": _build_provenance(data).get("browser_trace_ids", []),
            "tool_coverage_gaps": data.get("tool_coverage_gaps", []),
        }),
        "limitations": [
            "结论仅基于任务时间窗内保存的规范化证据。",
            "未达到确定性判定器门槛的项目保留为不确定或负面证据。",
        ],
        "remediation": _scrub(data.get("remediation", [])),
        "provenance": _build_provenance(data),
    }
    canonical = json.dumps(document, ensure_ascii=False, sort_keys=True, separators=(",", ":"), default=str)
    document["integrity_hash"] = hashlib.sha256(canonical.encode("utf-8")).hexdigest()
    return document


class ReportExporter:
    def __init__(self, document: dict[str, Any]):
        self.document = _scrub(document)
        self.document.setdefault("integrity_hash", self._digest())

    def render(self, fmt: ReportFormat) -> bytes:
        if fmt == "json":
            return json.dumps(self.document, ensure_ascii=False, indent=2, default=str).encode("utf-8")
        if fmt == "md":
            return self._markdown().encode("utf-8")
        if fmt == "html":
            return self._html().encode("utf-8")
        if fmt == "docx":
            return self._docx()
        if fmt == "pdf":
            return self._pdf()
        raise ValueError(f"unsupported report format: {fmt}")

    def export(self, output_dir: str | Path, *, formats: Iterable[ReportFormat] = ("md", "html", "json", "pdf", "docx")) -> dict[str, Any]:
        directory = Path(output_dir)
        directory.mkdir(parents=True, exist_ok=True)
        paths: dict[str, str] = {}
        for fmt in formats:
            safe_fmt = str(fmt).lower()
            payload = self.render(safe_fmt)  # type: ignore[arg-type]
            destination = directory / f"report.{safe_fmt}"
            destination.write_bytes(payload)
            paths[safe_fmt] = str(destination)
        return {"document": self.document, "paths": paths, "integrity_hash": self.document["integrity_hash"]}

    def _digest(self) -> str:
        payload = dict(self.document)
        payload.pop("integrity_hash", None)
        return hashlib.sha256(json.dumps(payload, ensure_ascii=False, sort_keys=True, default=str).encode("utf-8")).hexdigest()

    def _markdown(self) -> str:
        lines = [
            f"# SDIT Task Report `{self.document.get('task_id', '')}`",
            "",
            f"- Schema: `{self.document.get('schema_version', '')}`",
            f"- Generated: `{self.document.get('generated_at', '')}`",
            f"- Integrity: `{self.document.get('integrity_hash', '')}`",
            "",
        ]
        for section in ("management_summary", "technical_findings", "attack_path", "evidence", "failure_path", "coverage", "limitations", "remediation", "provenance"):
            lines.extend([f"## {section}", "", "```json", json.dumps(self.document.get(section, {}), ensure_ascii=False, indent=2, default=str), "```", ""])
        return "\n".join(lines)

    def _html(self) -> str:
        title = html.escape(f"SDIT Task Report {self.document.get('task_id', '')}")
        body = html.escape(self._markdown())
        return f"<!doctype html><html lang=\"zh-CN\"><head><meta charset=\"utf-8\"><title>{title}</title><style>body{{font-family:system-ui,sans-serif;max-width:1100px;margin:2rem auto;line-height:1.5}}pre{{white-space:pre-wrap;background:#f4f4f4;padding:1rem;border-radius:.5rem}}</style></head><body><pre>{body}</pre></body></html>"

    def _docx(self) -> bytes:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - dependency gate handles clean installs
            raise RuntimeError("python-docx is required for DOCX export") from exc
        document = Document()
        document.add_heading(f"SDIT Task Report {self.document.get('task_id', '')}", level=1)
        document.add_paragraph(f"Integrity: {self.document.get('integrity_hash', '')}")
        for section in ("management_summary", "technical_findings", "attack_path", "evidence", "failure_path", "coverage", "limitations", "remediation", "provenance"):
            document.add_heading(section, level=2)
            document.add_paragraph(json.dumps(self.document.get(section, {}), ensure_ascii=False, indent=2, default=str))
        from io import BytesIO

        output = BytesIO()
        document.save(output)
        return output.getvalue()

    def _pdf(self) -> bytes:
        # Minimal self-contained PDF writer; non-ASCII glyphs are replaced while
        # IDs, hashes and structured facts stay readable in every PDF viewer.
        text = self._markdown().encode("ascii", "replace").decode("ascii")
        lines = []
        for raw_line in text.splitlines():
            lines.extend(textwrap.wrap(raw_line, width=105) or [""])
        pages = [lines[index:index + 48] for index in range(0, len(lines), 48)] or [[]]
        objects: list[bytes] = []
        page_ids: list[int] = []
        objects.extend([b"<< /Type /Catalog /Pages 2 0 R >>", b""])
        # Page tree and font are filled after page/content object allocation.
        for page in pages:
            content = ["BT", "/F1 9 Tf", "50 760 Td"]
            for line in page:
                escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
                content.append(f"({escaped}) Tj 0 -14 Td")
            content.append("ET")
            stream = "\n".join(content).encode("latin-1", "replace")
            content_id = len(objects) + 1
            objects.append(f"<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream")
            page_id = len(objects) + 1
            objects.append(b"")
            page_ids.append(page_id)
        font_id = len(objects) + 1
        objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
        for page_id, content_id in zip(page_ids, range(3, 3 + 2 * len(page_ids), 2), strict=True):
            objects[page_id - 1] = (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {content_id} 0 R >>"
            ).encode()
        kids = " ".join(f"{item} 0 R" for item in page_ids)
        objects[1] = f"<< /Type /Pages /Kids [{kids}] /Count {len(page_ids)} >>".encode()
        output = bytearray(b"%PDF-1.4\n")
        offsets = [0]
        for index, obj in enumerate(objects, 1):
            offsets.append(len(output))
            output.extend(f"{index} 0 obj\n".encode())
            output.extend(obj)
            output.extend(b"\nendobj\n")
        xref = len(output)
        output.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode())
        for offset in offsets[1:]:
            output.extend(f"{offset:010d} 00000 n \n".encode())
        output.extend(f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode())
        return bytes(output)


__all__ = ["ReportFormat", "build_report_document", "ReportExporter"]
